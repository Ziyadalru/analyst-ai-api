import io
import re
import pandas as pd
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.concurrency import run_in_threadpool
import pdfplumber
from docx import Document as DocxDocument

from core.session_store import save_session
from core.numpy_encoder import numpy_safe
from modules.column_detector import detect_columns_smart, get_data_quality

router = APIRouter()


def _header_score(df: pd.DataFrame) -> float:
    """Score a parsed DataFrame — higher = better header row was chosen.
    Penalises Unnamed columns and high null rates."""
    unnamed = sum(1 for c in df.columns if str(c).startswith('Unnamed') or str(c).startswith('nan'))
    null_rate = df.isnull().sum().sum() / max(df.size, 1)
    col_penalty = unnamed / max(len(df.columns), 1)
    return -(col_penalty * 2 + null_rate)


def _smart_parse_sheet(xl: pd.ExcelFile, sheet: str) -> pd.DataFrame:
    """Parse an Excel sheet, auto-detecting the real header row.

    Academic / coursework files commonly have:
    - Title rows at the top (merged cells, module name, etc.)
    - Blank rows before data starts
    - Multi-row sub-headers

    Strategy: try header rows 0–9, pick whichever gives the fewest
    Unnamed columns and lowest null rate, then clean up.
    """
    # Step 1: fast scan with nrows=50 to find the best header row
    best_header_row, best_score = 0, -999
    for header_row in range(10):
        try:
            df = xl.parse(sheet, header=header_row, nrows=50)
            df = df.dropna(how='all').dropna(axis=1, how='all')
            if df.empty or len(df.columns) < 2:
                continue
            score = _header_score(df)
            if score > best_score:
                best_score, best_header_row = score, header_row
        except Exception:
            continue

    # Step 2: one full parse with the winning header row
    try:
        best_df = xl.parse(sheet, header=best_header_row)
    except Exception:
        best_df = xl.parse(sheet)

    # Forward-fill merged header cells (appear as NaN after the merge)
    best_df.columns = pd.Series(best_df.columns).ffill().tolist()

    # Clean column names
    best_df.columns = [
        re.sub(r'\s+', ' ', str(c).strip()).replace('\n', ' ')
        for c in best_df.columns
    ]

    # Drop rows that are entirely NaN or look like sub-header repeat rows
    best_df = best_df.dropna(how='all').reset_index(drop=True)

    return best_df


def _best_sheet(xl: pd.ExcelFile) -> str:
    """Return the sheet with the most usable data cells."""
    best, best_score = xl.sheet_names[0], -1
    for name in xl.sheet_names:
        try:
            tmp = xl.parse(name, nrows=200, header=None)
            score = tmp.size - tmp.isnull().sum().sum()
            if score > best_score:
                best_score, best = score, name
        except Exception:
            pass
    return best


def _parse_pdf(contents: bytes) -> pd.DataFrame:
    """Extract the largest table from a PDF. Falls back to text-as-rows if no tables found."""
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        # Try to find tables across all pages
        all_tables = []
        for page in pdf.pages:
            tables = page.extract_tables()
            for t in tables:
                if t and len(t) > 1:
                    all_tables.append(t)

        if all_tables:
            # Pick the largest table
            best = max(all_tables, key=lambda t: len(t) * len(t[0]))
            headers = [str(h).strip() if h else f'Col_{i}' for i, h in enumerate(best[0])]
            rows = [[str(c).strip() if c else '' for c in row] for row in best[1:]]
            df = pd.DataFrame(rows, columns=headers)
            # Try to convert numeric columns
            for col in df.columns:
                try:
                    df[col] = pd.to_numeric(df[col].str.replace(',', ''), errors='ignore')
                except Exception:
                    pass
            return df.replace('', pd.NA).dropna(how='all').dropna(axis=1, how='all')

        # No tables found — extract text lines as a single-column dataset
        text_rows = []
        for page in pdf.pages:
            text = page.extract_text() or ''
            text_rows.extend([l.strip() for l in text.split('\n') if l.strip()])
        if not text_rows:
            raise ValueError("No extractable content found in PDF.")
        return pd.DataFrame({'text': text_rows})


def _parse_word(contents: bytes) -> pd.DataFrame:
    """Extract the largest table from a Word doc. Falls back to paragraph text if no tables."""
    doc = DocxDocument(io.BytesIO(contents))

    if doc.tables:
        # Pick largest table
        best = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
        data = [[cell.text.strip() for cell in row.cells] for row in best.rows]
        if len(data) > 1:
            headers = data[0]
            # Deduplicate headers
            seen: dict[str, int] = {}
            clean_headers = []
            for h in headers:
                h = h or 'Col'
                if h in seen:
                    seen[h] += 1
                    clean_headers.append(f'{h}_{seen[h]}')
                else:
                    seen[h] = 0
                    clean_headers.append(h)
            df = pd.DataFrame(data[1:], columns=clean_headers)
            for col in df.columns:
                try:
                    df[col] = pd.to_numeric(df[col].str.replace(',', ''), errors='ignore')
                except Exception:
                    pass
            return df.replace('', pd.NA).dropna(how='all').dropna(axis=1, how='all')

    # No tables — use paragraphs as text rows
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError("No extractable content found in Word document.")
    return pd.DataFrame({'text': paragraphs})


def _parse_excel(contents: bytes, sheet_name: str | None) -> tuple[pd.DataFrame, list[str], str]:
    xl = pd.ExcelFile(io.BytesIO(contents))
    sheets   = xl.sheet_names
    selected = sheet_name if (sheet_name and sheet_name in sheets) else _best_sheet(xl)
    df       = _smart_parse_sheet(xl, selected)
    return df, sheets, selected


@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    sheet_name: str = Form(None),
):
    fname    = (file.filename or "").lower()
    is_excel = any(fname.endswith(ext) for ext in ('.xlsx', '.xls', '.xlsm'))
    is_csv   = fname.endswith('.csv')
    is_pdf   = fname.endswith('.pdf')
    is_word  = any(fname.endswith(ext) for ext in ('.docx', '.doc'))

    if not any([is_excel, is_csv, is_pdf, is_word]):
        raise HTTPException(status_code=400, detail="Supported formats: CSV, Excel (.xlsx/.xls), PDF, Word (.docx).")

    contents = await file.read()
    sheets: list[str] = []
    active_sheet: str | None = None

    try:
        if is_csv:
            df = pd.read_csv(io.BytesIO(contents), encoding="unicode_escape")
        elif is_excel:
            df, sheets, active_sheet = _parse_excel(contents, sheet_name)
        elif is_pdf:
            df = _parse_pdf(contents)
        elif is_word:
            df = _parse_word(contents)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {e}")

    if df.empty:
        raise HTTPException(status_code=422, detail="File is empty.")

    # Final clean — drop fully empty rows/cols, reset index
    df = df.dropna(how='all').dropna(axis=1, how='all').reset_index(drop=True)

    if df.empty:
        raise HTTPException(status_code=422, detail="No usable data found. Try selecting a different sheet.")

    cols, ai_meta = await run_in_threadpool(detect_columns_smart, df)
    session_id   = save_session(df, cols)

    suggested = ai_meta['suggested_analyses']
    if "Executive Dashboard" not in suggested:
        suggested = ["Executive Dashboard"] + suggested

    quality = get_data_quality(df)

    return {
        "session_id":         session_id,
        "quality":            numpy_safe(quality),
        "detected_columns":   {k: v for k, v in cols.items() if v is not None},
        "suggested_analyses": suggested,
        "column_names":       df.columns.tolist(),
        "row_count":          len(df),
        "sheets":             sheets,
        "active_sheet":       active_sheet,
        "file_type":          "excel" if is_excel else "pdf" if is_pdf else "word" if is_word else "csv",
        "data_type":          ai_meta['data_type'],
        "data_description":   ai_meta['description'],
        "data_notes":         ai_meta['notes'],
    }
